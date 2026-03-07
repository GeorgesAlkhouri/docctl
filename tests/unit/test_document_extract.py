from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from docctl import document_extract
from docctl.errors import EmptyExtractedTextError, InputPathNotFoundError
from docctl.models import TextUnit


def test_extract_text_units_splits_paragraphs(tmp_path: Path) -> None:
    text_path = tmp_path / "notes.txt"
    text_path.write_text("First paragraph.\n\nSecond paragraph.", encoding="utf-8")

    units = document_extract.extract_document_units(text_path)

    assert units == [
        TextUnit(text="First paragraph."),
        TextUnit(text="Second paragraph."),
    ]


def test_extract_markdown_units_preserve_order(tmp_path: Path) -> None:
    markdown_path = tmp_path / "notes.md"
    markdown_path.write_text("# Header\n\nBody paragraph.", encoding="utf-8")

    units = document_extract.extract_document_units(markdown_path)

    assert [unit.text for unit in units] == ["# Header", "Body paragraph."]


def test_extract_docx_units_reads_non_empty_paragraphs(tmp_path: Path) -> None:
    docx_path = tmp_path / "notes.docx"
    document = Document()
    document.add_paragraph("First docx paragraph.")
    document.add_paragraph("")
    document.add_paragraph("Second docx paragraph.")
    document.save(str(docx_path))

    units = document_extract.extract_document_units(docx_path)

    assert units == [
        TextUnit(text="First docx paragraph."),
        TextUnit(text="Second docx paragraph."),
    ]


def test_extract_docx_units_raises_read_error(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    docx_path = tmp_path / "broken.docx"
    docx_path.write_bytes(b"not-a-docx")
    monkeypatch.setattr(
        document_extract,
        "DocxDocument",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(ValueError("bad")),
    )

    with pytest.raises(document_extract.DocumentReadError):
        document_extract.extract_document_units(docx_path)


def test_extract_docx_units_raises_when_no_text(tmp_path: Path) -> None:
    docx_path = tmp_path / "empty.docx"
    document = Document()
    document.add_paragraph("")
    document.save(str(docx_path))

    with pytest.raises(EmptyExtractedTextError):
        document_extract.extract_document_units(docx_path)


def test_extract_document_units_raises_for_unsupported_extension(tmp_path: Path) -> None:
    path = tmp_path / "notes.csv"
    path.write_text("a,b\n1,2", encoding="utf-8")

    with pytest.raises(InputPathNotFoundError):
        document_extract.extract_document_units(path)


def test_extract_document_units_raises_for_empty_text_file(tmp_path: Path) -> None:
    path = tmp_path / "empty.txt"
    path.write_text(" \n\t", encoding="utf-8")

    with pytest.raises(EmptyExtractedTextError):
        document_extract.extract_document_units(path)


def test_extract_text_units_raises_read_error(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    path = tmp_path / "notes.txt"
    path.write_text("content", encoding="utf-8")
    monkeypatch.setattr(
        Path,
        "read_text",
        lambda self, encoding=None: (_ for _ in ()).throw(OSError("denied")),
    )

    with pytest.raises(document_extract.DocumentReadError):
        document_extract.extract_document_units(path)
